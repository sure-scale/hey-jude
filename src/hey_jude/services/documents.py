import base64
import binascii
from dataclasses import asdict, dataclass
from email import policy
from email.parser import BytesParser
from html.parser import HTMLParser
from io import BytesIO
import re
import zipfile
from typing import Any, Literal


UnreadableAction = Literal["reject", "warn", "skip"]


@dataclass
class DocumentWarning:
    filename: str | None
    media_type: str | None
    reason: str
    action: str

    def to_dict(self) -> dict[str, str | None]:
        return asdict(self)


@dataclass
class DocumentContentResult:
    text: str
    warnings: list[DocumentWarning]


class DocumentProcessingError(ValueError):
    pass


class _HTMLTextParser(HTMLParser):
    def __init__(self):
        super().__init__()
        self._parts: list[str] = []

    def handle_data(self, data: str):
        text = data.strip()
        if text:
            self._parts.append(text)

    def text(self) -> str:
        return "\n".join(self._parts)


def _html_to_text(raw: bytes) -> str:
    try:
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(raw, "html.parser")
        return soup.get_text("\n", strip=True)
    except Exception:
        parser = _HTMLTextParser()
        parser.feed(raw.decode("utf-8", errors="replace"))
        return parser.text()


def _docx_to_text(raw: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise DocumentProcessingError("DOCX support is not installed") from exc

    document = Document(BytesIO(raw))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def _pdf_to_text(raw: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise DocumentProcessingError("PDF support is not installed") from exc

    reader = PdfReader(BytesIO(raw))
    parts = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text.strip())
    return "\n".join(parts)


def _rtf_to_text(raw: bytes) -> str:
    text = raw.decode("utf-8", errors="replace")
    text = re.sub(r"\\'[0-9a-fA-F]{2}", " ", text)
    text = re.sub(r"\\[a-zA-Z]+-?\d* ?", " ", text)
    text = text.replace("{", " ").replace("}", " ")
    return re.sub(r"\s+", " ", text).strip()


def _email_to_text(raw: bytes) -> str:
    message = BytesParser(policy=policy.default).parsebytes(raw)
    parts = []
    subject = message.get("subject")
    if subject:
        parts.append(f"Subject: {subject}")

    plain_parts = []
    html_parts = []
    if message.is_multipart():
        for part in message.walk():
            disposition = part.get_content_disposition()
            if disposition == "attachment":
                continue
            content_type = part.get_content_type()
            payload = part.get_payload(decode=True)
            if not payload:
                continue
            if content_type == "text/plain":
                plain_parts.append(
                    payload.decode(
                        part.get_content_charset() or "utf-8",
                        errors="replace",
                    )
                )
            elif content_type == "text/html":
                html_parts.append(_html_to_text(payload))
    else:
        payload = message.get_payload(decode=True)
        if payload:
            if message.get_content_type() == "text/html":
                html_parts.append(_html_to_text(payload))
            else:
                plain_parts.append(
                    payload.decode(
                        message.get_content_charset() or "utf-8",
                        errors="replace",
                    )
                )

    parts.extend(part.strip() for part in plain_parts if part.strip())
    if not plain_parts:
        parts.extend(part.strip() for part in html_parts if part.strip())
    return "\n".join(parts)


def extract_document_text(
    raw: bytes,
    media_type: str | None = None,
    filename: str | None = None,
) -> str:
    media_type = (media_type or "").split(";")[0].strip().lower()
    suffix = ""
    if filename and "." in filename:
        suffix = filename.rsplit(".", 1)[1].lower()

    if media_type.startswith("text/plain") or suffix in {"txt", "md"}:
        text = raw.decode("utf-8", errors="replace")
    elif media_type == "text/html" or suffix in {"html", "htm"}:
        text = _html_to_text(raw)
    elif media_type in {"message/rfc822", "application/eml"} or suffix == "eml":
        text = _email_to_text(raw)
    elif media_type in {"application/rtf", "text/rtf"} or suffix == "rtf":
        text = _rtf_to_text(raw)
    elif media_type == "application/pdf" or suffix == "pdf":
        text = _pdf_to_text(raw)
    elif (
        media_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or suffix == "docx"
    ):
        text = _docx_to_text(raw)
    else:
        kind = media_type or (f".{suffix}" if suffix else "attachment")
        raise DocumentProcessingError(f"{kind} is not readable as text")

    if not text.strip():
        kind = media_type or (f".{suffix}" if suffix else "attachment")
        raise DocumentProcessingError(f"{kind} has no readable text layer")
    return text.strip()


def _decode_base64_payload(value: str) -> tuple[bytes, str | None]:
    try:
        if value.startswith("data:"):
            header, _, payload = value.partition(",")
            if not payload:
                raise DocumentProcessingError("data URL is missing base64 payload")
            media_type = header[5:].split(";")[0] or None
            return base64.b64decode(payload, validate=True), media_type
        return base64.b64decode(value, validate=True), None
    except binascii.Error as exc:
        raise DocumentProcessingError("attachment has invalid base64 data") from exc


def _part_attachment(part: dict[str, Any]) -> tuple[bytes, str | None, str | None] | None:
    filename = part.get("filename") or part.get("name")
    media_type = part.get("media_type") or part.get("mime_type")

    if isinstance(part.get("file_data"), str):
        raw, detected_type = _decode_base64_payload(part["file_data"])
        return raw, media_type or detected_type, filename

    if isinstance(part.get("data"), str):
        raw, detected_type = _decode_base64_payload(part["data"])
        return raw, media_type or detected_type, filename

    source = part.get("source")
    if isinstance(source, dict):
        source_type = source.get("type")
        source_media_type = source.get("media_type") or media_type
        if source_type == "base64" and isinstance(source.get("data"), str):
            raw, detected_type = _decode_base64_payload(source["data"])
            return raw, source_media_type or detected_type, filename

    image_url = part.get("image_url")
    if isinstance(image_url, dict) and isinstance(image_url.get("url"), str):
        raw, detected_type = _decode_base64_payload(image_url["url"])
        return raw, media_type or detected_type or "image", filename

    inline_data = part.get("inline_data")
    if isinstance(inline_data, dict) and isinstance(inline_data.get("data"), str):
        raw, detected_type = _decode_base64_payload(inline_data["data"])
        return raw, inline_data.get("mime_type") or detected_type, filename

    return None


def _handle_unreadable(
    warning: DocumentWarning,
    action: UnreadableAction,
) -> str | None:
    if action == "reject":
        raise DocumentProcessingError(warning.reason)
    if action == "warn":
        return (
            "[Attachment omitted: content is not readable as text. "
            f"Reason: {warning.reason}]"
        )
    return None


def _part_to_text(
    part: Any,
    warnings: list[DocumentWarning],
    action: UnreadableAction,
) -> str | None:
    if isinstance(part, str):
        return part
    if not isinstance(part, dict):
        return None

    text = part.get("text")
    if isinstance(text, str):
        return text

    try:
        attachment = _part_attachment(part)
        if attachment is None:
            return None
        raw, media_type, filename = attachment
        return extract_document_text(raw, media_type=media_type, filename=filename)
    except (DocumentProcessingError, zipfile.BadZipFile, OSError) as exc:
        media_type = locals().get("media_type")
        filename = locals().get("filename")
        reason = str(exc)
        if "not readable as text" not in reason and "readable text layer" not in reason:
            kind = media_type or "attachment"
            reason = f"{kind} is not readable as text: {reason}"
        warning = DocumentWarning(
            filename=filename,
            media_type=media_type,
            reason=reason,
            action=action,
        )
        warnings.append(warning)
        return _handle_unreadable(warning, action)


def content_to_text(content: Any, settings) -> DocumentContentResult:
    action: UnreadableAction = settings.document_unreadable_action
    warnings: list[DocumentWarning] = []

    if isinstance(content, str):
        return DocumentContentResult(text=content, warnings=warnings)

    if isinstance(content, dict):
        text = _part_to_text(content, warnings, action)
        return DocumentContentResult(text=text or "", warnings=warnings)

    if isinstance(content, list):
        parts = []
        for part in content:
            text = _part_to_text(part, warnings, action)
            if text:
                parts.append(text)
        return DocumentContentResult(text="\n".join(parts), warnings=warnings)

    return DocumentContentResult(text="", warnings=warnings)
